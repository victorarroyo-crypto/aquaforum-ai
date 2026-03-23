"""
AquaForum AI — Professional Word Document Export Service.

Generates executive-quality .docx reports with:
  - Cover page with branding, date, topic, panelist roster
  - Introduction / context section
  - Auto-generated Table of Contents (Word field codes)
  - Per-round structured analysis (question, expert highlights, synthesis, summary)
  - Final summary with conclusions and recommendations
  - Annex: full literal transcript organised by round
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------
TEAL = RGBColor(0x14, 0xB8, 0xA6)
DARK_TEAL = RGBColor(0x0D, 0x94, 0x88)
DARK_GRAY = RGBColor(0x37, 0x41, 0x51)
MID_GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_GRAY = RGBColor(0x9C, 0xA3, 0xAF)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
ACCENT_BLUE = RGBColor(0x38, 0xBD, 0xF8)


# ---------------------------------------------------------------------------
# Spanish month names (avoid locale issues)
# ---------------------------------------------------------------------------
_MESES = [
    "", "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
]


def _fecha_es(dt: datetime) -> str:
    return f"{dt.day} de {_MESES[dt.month]} de {dt.year}"


# ---------------------------------------------------------------------------
# Message-type labels (Spanish)
# ---------------------------------------------------------------------------
_TYPE_LABELS: dict[str, str] = {
    "moderation": "Moderacion",
    "statement": "Declaracion",
    "challenge": "Interpelacion",
    "response": "Respuesta",
    "round_summary": "Resumen de Ronda",
    "summary": "Resumen Final",
    "analysis": "Analisis",
    "integration": "Sintesis Estrategica",
}


# ---------------------------------------------------------------------------
# Markdown-stripping helper
# ---------------------------------------------------------------------------
def _clean_md(text: str) -> str:
    """Remove lightweight Markdown formatting that looks odd in Word."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)   # bold markers
    text = re.sub(r"\*([^*]+)\*", r"\1", text)        # italic markers
    text = re.sub(r"#{1,4}\s*", "", text)              # heading markers
    text = re.sub(r"\[CHALLENGE:[^\]]*\]", "", text)   # challenge tags
    return text.strip()


def _extract_bullets(text: str) -> list[str]:
    """Pull out Markdown-style bullet lines from a block of text."""
    bullets: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(("- ", "* ", "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            clean = re.sub(r"^[-*\d.]+\s*", "", stripped)
            if clean:
                bullets.append(clean)
    return bullets


# ---------------------------------------------------------------------------
# Low-level OxmlElement helpers
# ---------------------------------------------------------------------------
def _add_page_number(section):
    """Insert a centered page-number field in the footer."""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(8)
    p.style.font.color.rgb = MID_GRAY

    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _insert_toc_field(doc: Document):
    """Insert a Word TOC field code so the user can update it on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    run._r.append(instr)

    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_separate)

    # Placeholder text shown before user updates TOC
    placeholder_run = p.add_run(
        "[Haga clic derecho aqui y seleccione 'Actualizar campo' para generar el indice]"
    )
    placeholder_run.font.color.rgb = MID_GRAY
    placeholder_run.font.size = Pt(10)
    placeholder_run.italic = True

    end_run = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_char_end)


def _set_cell_shading(cell, hex_color: str):
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# Style configuration
# ---------------------------------------------------------------------------
def _configure_styles(doc: Document):
    """Set up professional typography for the document."""
    # Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = DARK_GRAY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Headings
    for level, (size, color, bold) in {
        1: (20, DARK_TEAL, True),
        2: (15, TEAL, True),
        3: (12, DARK_GRAY, True),
    }.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = bold
        style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
        style.paragraph_format.space_after = Pt(8)

    # Margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------
def _build_cover(doc: Document, session: dict, now: datetime):
    """Full-bleed cover page."""
    config = session.get("config", {})
    panelists = config.get("panelists", [])

    # Vertical spacing to center content visually
    for _ in range(5):
        doc.add_paragraph("")

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("AquaForum AI")
    run.bold = True
    run.font.size = Pt(38)
    run.font.color.rgb = TEAL

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Informe Ejecutivo")
    run.font.size = Pt(22)
    run.font.color.rgb = MID_GRAY

    doc.add_paragraph("")

    # Decorative rule
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_rule.add_run("_" * 50)
    run.font.color.rgb = TEAL
    run.font.size = Pt(10)

    doc.add_paragraph("")

    # Topic
    p_topic = doc.add_paragraph()
    p_topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_topic.add_run(session["topic"])
    run.font.size = Pt(16)
    run.italic = True
    run.font.color.rgb = DARK_GRAY

    for _ in range(2):
        doc.add_paragraph("")

    # Date
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_date.add_run(_fecha_es(now))
    run.font.size = Pt(12)
    run.font.color.rgb = MID_GRAY

    doc.add_paragraph("")

    # Panelist roster
    if panelists:
        p_pan_title = doc.add_paragraph()
        p_pan_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_pan_title.add_run("Panel de Expertos")
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = DARK_TEAL

        for pan in panelists:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            name_run = p.add_run(f"{pan['name']}")
            name_run.bold = True
            name_run.font.size = Pt(11)
            name_run.font.color.rgb = DARK_GRAY
            role_run = p.add_run(f"  —  {pan['role']}")
            role_run.font.size = Pt(10)
            role_run.font.color.rgb = MID_GRAY

    for _ in range(3):
        doc.add_paragraph("")

    # Footer line
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run(
        f"Generado automaticamente por AquaForum AI  |  "
        f"Rondas: {session.get('max_rounds', 3)}  |  "
        f"Motor: Claude + LangGraph"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = LIGHT_GRAY

    doc.add_page_break()


def _build_introduction(doc: Document, session: dict):
    """Context section: why this forum exists."""
    doc.add_heading("Introduccion", level=1)

    paras = [
        (
            "La inteligencia artificial esta transformando cada sector de la economia global "
            "a una velocidad sin precedentes. Desde la automatizacion de procesos industriales "
            "hasta la toma de decisiones basada en datos en tiempo real, las organizaciones "
            "de todos los ambitos estan redefiniendo sus modelos operativos. El sector del agua, "
            "sin embargo, enfrenta un conjunto de desafios unicos que hacen de esta transformacion "
            "algo particularmente urgente y complejo."
        ),
        (
            "La infraestructura hidrica mundial envejece de manera acelerada: se estima que un 30% "
            "del agua tratada se pierde por fugas en las redes de distribucion (IWA, 2024). "
            "El cambio climatico esta alterando los patrones de disponibilidad de recursos hidricos, "
            "generando episodios mas frecuentes de sequia e inundaciones. La presion regulatoria "
            "se intensifica con nuevas directivas europeas sobre aguas residuales, reutilizacion "
            "y ciberseguridad de infraestructuras criticas (NIS2). Y todo esto ocurre en un contexto "
            "donde la demanda de agua crece impulsada por la urbanizacion y el desarrollo economico."
        ),
        (
            "El escenario descrito en ai-2027.com sugiere que la aceleracion tecnologica "
            "hacia sistemas de IA cada vez mas capaces podria ser inminente. Este contexto "
            "de urgencia motivado por ai-2027.com actua como catalizador para la creacion "
            "de AquaForum AI: un espacio donde expertos en IA y profesionales del sector "
            "hidrico debaten, desde perspectivas complementarias, las implicaciones practicas "
            "de esta transformacion."
        ),
        (
            "El presente informe recoge las deliberaciones del foro, organizadas por rondas "
            "de debate, con analisis de expertos, sintesis estrategicas y la transcripcion "
            "completa de todas las intervenciones. Su objetivo es proporcionar a los tomadores "
            "de decisiones del sector del agua una vision integral y accionable sobre como "
            "prepararse para la era de la IA avanzada."
        ),
    ]

    for text in paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(10)

    doc.add_page_break()


def _build_toc(doc: Document):
    """Table of Contents page using Word field codes."""
    doc.add_heading("Indice de Contenidos", level=1)

    p_note = doc.add_paragraph()
    run = p_note.add_run(
        "Este indice se actualiza automaticamente en Microsoft Word: "
        "haga clic derecho sobre el indice y seleccione 'Actualizar campo'."
    )
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = MID_GRAY

    doc.add_paragraph("")
    _insert_toc_field(doc)
    doc.add_page_break()


def _build_round_section(doc: Document, round_num: int, messages: list[dict]):
    """Structured analysis for a single debate round."""
    round_msgs = [m for m in messages if m["round_number"] == round_num]
    if not round_msgs:
        return

    doc.add_heading(f"Ronda {round_num}", level=1)

    # --- Moderator opening / guiding question ---
    openings = [
        m for m in round_msgs
        if m["message_type"] == "moderation" and m.get("metadata", {}).get("action") == "open"
    ]
    if openings:
        doc.add_heading("Pregunta orientadora", level=2)
        doc.add_paragraph(_clean_md(openings[0]["content"]))

    # --- Expert analysis highlights ---
    analyses = [m for m in round_msgs if m["message_type"] == "analysis"]
    if analyses:
        doc.add_heading("Analisis de Expertos", level=2)
        for a in analyses:
            # Expert name as sub-heading
            doc.add_heading(a["agent_name"], level=3)

            content = _clean_md(a["content"])
            bullets = _extract_bullets(content)

            if bullets:
                for b in bullets[:6]:  # cap at 6 bullets per expert
                    p = doc.add_paragraph(style="List Bullet")
                    p.text = _clean_md(b)
            else:
                # Fallback: first 3 sentences as highlight
                sentences = re.split(r"(?<=[.!?])\s+", content)
                highlight = " ".join(sentences[:3])
                doc.add_paragraph(highlight)

    # --- Integrator synthesis ---
    integrations = [m for m in round_msgs if m["message_type"] == "integration"]
    if integrations:
        doc.add_heading("Sintesis Estrategica", level=2)
        content = _clean_md(integrations[0]["content"])
        bullets = _extract_bullets(content)
        if bullets:
            # First paragraph as narrative intro
            non_bullet_lines = [
                line.strip() for line in content.splitlines()
                if line.strip() and not line.strip().startswith(("- ", "* ", "1.", "2.", "3."))
            ]
            if non_bullet_lines:
                doc.add_paragraph(non_bullet_lines[0])
            for b in bullets[:8]:
                p = doc.add_paragraph(style="List Bullet")
                p.text = _clean_md(b)
        else:
            doc.add_paragraph(content)

    # --- Moderator round summary ---
    summaries = [m for m in round_msgs if m["message_type"] == "round_summary"]
    if summaries:
        doc.add_heading("Resumen del Moderador", level=2)
        content = _clean_md(summaries[0]["content"])
        bullets = _extract_bullets(content)
        if bullets:
            non_bullet = [
                line.strip() for line in content.splitlines()
                if line.strip() and not line.strip().startswith(("- ", "* ", "1.", "2.", "3."))
            ]
            if non_bullet:
                doc.add_paragraph(non_bullet[0])
            for b in bullets[:6]:
                p = doc.add_paragraph(style="List Bullet")
                p.text = _clean_md(b)
        else:
            doc.add_paragraph(content)

    doc.add_page_break()


def _build_final_summary(doc: Document, messages: list[dict]):
    """Conclusions and recommendations from the final summary."""
    finals = [m for m in messages if m["message_type"] == "summary"]
    if not finals:
        return

    doc.add_heading("Conclusiones y Recomendaciones", level=1)

    content = _clean_md(finals[-1]["content"])

    # Try to split into a narrative intro + bullet recommendations
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]

    for para_text in paragraphs:
        bullets = _extract_bullets(para_text)
        if bullets:
            for b in bullets:
                p = doc.add_paragraph(style="List Bullet")
                p.text = _clean_md(b)
        else:
            lines = para_text.strip().splitlines()
            combined = " ".join(line.strip() for line in lines if line.strip())
            if combined:
                doc.add_paragraph(combined)

    doc.add_page_break()


def _build_transcript_annex(doc: Document, messages: list[dict]):
    """Full literal transcript organised by round."""
    doc.add_heading("Anexo: Transcripcion Completa", level=1)

    p_intro = doc.add_paragraph()
    run = p_intro.add_run(
        "A continuacion se incluye la transcripcion literal e integra de todas las "
        "intervenciones del debate, organizada por rondas de discusion. Cada intervencion "
        "incluye el nombre del participante, su rol y el tipo de intervencion."
    )
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = MID_GRAY

    doc.add_paragraph("")

    current_round = 0
    for msg in messages:
        # Round heading
        if msg["round_number"] != current_round:
            current_round = msg["round_number"]
            doc.add_heading(f"Ronda {current_round}", level=2)

        type_label = _TYPE_LABELS.get(msg["message_type"], msg["message_type"])

        # Speaker header with role
        p_speaker = doc.add_paragraph()
        p_speaker.paragraph_format.space_before = Pt(10)
        p_speaker.paragraph_format.space_after = Pt(2)

        name_run = p_speaker.add_run(msg["agent_name"])
        name_run.bold = True
        name_run.font.size = Pt(10)
        name_run.font.color.rgb = DARK_TEAL

        role_run = p_speaker.add_run(f"  ({msg['agent_role']})")
        role_run.font.size = Pt(9)
        role_run.font.color.rgb = MID_GRAY

        type_run = p_speaker.add_run(f"  —  {type_label}")
        type_run.font.size = Pt(9)
        type_run.italic = True
        type_run.font.color.rgb = MID_GRAY

        # Message content
        content = _clean_md(msg["content"])
        p_content = doc.add_paragraph(content)
        p_content.paragraph_format.space_after = Pt(8)
        for run in p_content.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_GRAY


# ---------------------------------------------------------------------------
# Main export function
# ---------------------------------------------------------------------------
def generate_docx(session: dict, messages: list[dict]) -> io.BytesIO:
    """
    Build a professional Word document for an AquaForum AI session.

    Parameters
    ----------
    session : dict
        The session record from Supabase (must include 'topic', 'config', 'max_rounds', etc.)
    messages : list[dict]
        All forum messages for the session, ordered chronologically.

    Returns
    -------
    io.BytesIO
        In-memory buffer containing the .docx file, ready to stream.
    """
    doc = Document()
    now = datetime.now(timezone.utc)

    # 1. Configure styles and page setup
    _configure_styles(doc)

    # Enable distinct first-page header/footer (cover has no page number)
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    _add_page_number(section)

    # 2. Cover page
    _build_cover(doc, session, now)

    # 3. Introduction
    _build_introduction(doc, session)

    # 4. Table of Contents
    _build_toc(doc)

    # 5. Round-by-round analysis
    max_round = max((m["round_number"] for m in messages), default=0)
    for r in range(1, max_round + 1):
        _build_round_section(doc, r, messages)

    # 6. Final summary / conclusions
    _build_final_summary(doc, messages)

    # 7. Annex: full transcript
    _build_transcript_annex(doc, messages)

    # Serialise
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
