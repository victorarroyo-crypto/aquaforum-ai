from __future__ import annotations

import asyncio

from fastapi import APIRouter, BackgroundTasks, HTTPException

from app.models.schemas import (
    ExportResponse,
    ForumConfig,
    SessionState,
    StartForumResponse,
)
from app.services import supabase_client as db

router = APIRouter()

# Track running tasks to prevent concurrent cycles
_running_sessions: set[str] = set()


@router.post("/start", response_model=StartForumResponse)
async def start_forum(config: ForumConfig, background_tasks: BackgroundTasks):
    session_id = await db.create_session(
        topic=config.topic,
        config=config.model_dump(),
        max_rounds=config.max_rounds,
    )

    # Run the first cycle in the background
    _running_sessions.add(session_id)
    background_tasks.add_task(_run_cycle, session_id, config, 1)

    return StartForumResponse(session_id=session_id, status="running")


@router.post("/{session_id}/next-cycle", status_code=202)
async def next_cycle(session_id: str, background_tasks: BackgroundTasks):
    session = await db.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    if session_id in _running_sessions:
        raise HTTPException(status_code=409, detail="A cycle is already running for this session")

    if session["status"] == "completed":
        raise HTTPException(status_code=400, detail="Session already completed")

    next_round = session["current_round"] + 1
    if next_round > session["max_rounds"]:
        raise HTTPException(status_code=400, detail="All rounds completed")

    config = ForumConfig(**session["config"])
    _running_sessions.add(session_id)
    background_tasks.add_task(_run_cycle, session_id, config, next_round)

    return {"status": "started", "cycle": next_round}


@router.post("/{session_id}/stop")
async def stop_forum(session_id: str):
    session = await db.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    await db.update_session_status(session_id, "paused")
    _running_sessions.discard(session_id)
    return {"status": "paused"}


@router.get("/{session_id}/state", response_model=SessionState)
async def get_state(session_id: str):
    session = await db.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    messages = await db.get_messages(session_id)

    return SessionState(
        session_id=session["id"],
        topic=session["topic"],
        status=session["status"],
        current_round=session["current_round"],
        max_rounds=session["max_rounds"],
        config=session["config"],
        messages=messages,
    )


@router.post("/{session_id}/export", response_model=ExportResponse)
async def export_forum(session_id: str):
    session = await db.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    messages = await db.get_messages(session_id)

    # Build a Markdown report
    lines = [
        f"# AquaForum AI — Informe del Foro",
        f"## Tema: {session['topic']}",
        f"**Estado:** {session['status']} | **Rondas:** {session['current_round']}/{session['max_rounds']}",
        "",
        "---",
        "",
    ]

    current_round = 0
    for msg in messages:
        if msg["round_number"] != current_round:
            current_round = msg["round_number"]
            lines.append(f"## Ronda {current_round}")
            lines.append("")

        prefix = f"**{msg['agent_name']}** ({msg['agent_role']}) — _{msg['message_type']}_"
        lines.append(prefix)
        lines.append(msg["content"])
        lines.append("")

    content = "\n".join(lines)
    report_id = await db.save_report(session_id, "full", content)

    return ExportResponse(report_id=report_id, content=content)


@router.get("/{session_id}/export-docx")
async def export_docx(session_id: str):
    """Export professional Word report with cover, TOC, analyses, and transcript annex."""
    import io
    import re
    from datetime import datetime

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor, Cm
    from fastapi.responses import StreamingResponse

    session = await db.get_session(session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    messages = await db.get_messages(session_id)
    doc = Document()

    # --- Styles ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    def clean_md(text: str) -> str:
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"#{1,3}\s*", "", text)
        text = re.sub(r"\[CHALLENGE:[^\]]+\]", "", text)
        return text.strip()

    def add_styled_para(text: str, bold: bool = False, italic: bool = False, size: int = 11, color: tuple | None = None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        return p

    # ==========================================
    # COVER PAGE
    # ==========================================
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AQUAFORUM AI")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informe del Evento")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x71, 0x71, 0x7A)

    doc.add_paragraph("")

    topic_p = doc.add_paragraph()
    topic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = topic_p.add_run(session["topic"])
    run.font.size = Pt(14)
    run.italic = True

    for _ in range(4):
        doc.add_paragraph("")

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    now = datetime.utcnow().strftime("%d de %B de %Y")
    meta.add_run(f"Fecha: {now}\n").font.size = Pt(10)
    meta.add_run(f"Rondas: {session.get('max_rounds', 4)}\n").font.size = Pt(10)

    panelist_names = [p["name"] for p in session.get("config", {}).get("panelists", [])]
    if panelist_names:
        meta.add_run(f"Panelistas: {', '.join(panelist_names)}\n").font.size = Pt(10)

    meta.add_run("Generado por AquaForum AI · Claude · LangGraph").font.size = Pt(9)

    doc.add_page_break()

    # ==========================================
    # TABLE OF CONTENTS (manual)
    # ==========================================
    doc.add_heading("Índice", level=1)
    max_round = max((m["round_number"] for m in messages), default=0)

    toc_items = []
    for r in range(1, max_round + 1):
        toc_items.append(f"Ronda {r} — Resumen y Análisis")
    if any(m["message_type"] == "summary" for m in messages):
        toc_items.append("Resumen Final")
    toc_items.append("Anexo: Transcripción Completa")

    for i, item in enumerate(toc_items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")

    doc.add_page_break()

    # ==========================================
    # ROUND-BY-ROUND ANALYSIS
    # ==========================================
    for r in range(1, max_round + 1):
        doc.add_heading(f"Ronda {r}", level=1)

        # Round summary
        summaries = [m for m in messages if m["round_number"] == r and m["message_type"] == "round_summary"]
        if summaries:
            doc.add_heading("Resumen del Moderador", level=2)
            doc.add_paragraph(clean_md(summaries[0]["content"]))

        # Expert analyses
        analyses = [m for m in messages if m["round_number"] == r and m["message_type"] == "analysis"]
        if analyses:
            doc.add_heading("Análisis de Expertos", level=2)
            for a in analyses:
                p = doc.add_paragraph()
                run = p.add_run(a["agent_name"])
                run.bold = True
                run.font.color.rgb = RGBColor(0x14, 0xB8, 0xA6)
                run.font.size = Pt(12)

                # Extract key highlight (first 2 sentences)
                content = clean_md(a["content"])
                sentences = re.split(r"(?<=[.!?])\s+", content)
                highlight = " ".join(sentences[:3])
                doc.add_paragraph(highlight)

        # Integration
        integrations = [m for m in messages if m["round_number"] == r and m["message_type"] == "integration"]
        if integrations:
            doc.add_heading("Síntesis Estratégica", level=2)
            doc.add_paragraph(clean_md(integrations[0]["content"]))

        doc.add_page_break()

    # ==========================================
    # FINAL SUMMARY
    # ==========================================
    final = [m for m in messages if m["message_type"] == "summary"]
    if final:
        doc.add_heading("Resumen Final del Debate", level=1)
        doc.add_paragraph(clean_md(final[-1]["content"]))
        doc.add_page_break()

    # ==========================================
    # ANNEX: FULL TRANSCRIPT
    # ==========================================
    doc.add_heading("Anexo: Transcripción Completa", level=1)
    add_styled_para(
        "A continuación se incluye la transcripción literal de todas las intervenciones del debate.",
        italic=True, size=10, color=(0x71, 0x71, 0x7A),
    )
    doc.add_paragraph("")

    current_round = 0
    for msg in messages:
        if msg["round_number"] != current_round:
            current_round = msg["round_number"]
            doc.add_heading(f"Ronda {current_round}", level=2)

        # Skip system messages
        if msg["message_type"] in ("analysis", "integration"):
            continue

        # Speaker header
        type_label = {
            "moderation": "Moderador",
            "statement": "Declaración",
            "challenge": "Interpelación",
            "response": "Respuesta",
            "round_summary": "Resumen de Ronda",
            "summary": "Resumen Final",
        }.get(msg["message_type"], msg["message_type"])

        p = doc.add_paragraph()
        run = p.add_run(f"{msg['agent_name']} — {type_label}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x52, 0x52, 0x5B)

        content = clean_md(msg["content"])
        tp = doc.add_paragraph(content)
        tp.style.font.size = Pt(10)
        doc.add_paragraph("")  # spacer

    # ==========================================
    # Save and return
    # ==========================================
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"AquaForum_Informe_{session_id[:8]}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


async def _run_cycle(session_id: str, config: ForumConfig, round_number: int):
    """Run ALL debate rounds automatically."""
    try:
        await db.update_session_status(session_id, "running")

        from app.engine.runner import run_all_rounds

        await run_all_rounds(session_id, config)
        await db.update_session_status(session_id, "completed")

    except Exception as e:
        await db.update_session_status(session_id, "error")
        await db.add_message(
            session_id=session_id,
            agent_name="Sistema",
            agent_role="system",
            content=f"Error en la ejecución: {str(e)}",
            message_type="moderation",
            round_number=round_number,
            turn_number=0,
            metadata={"error": True},
        )
    finally:
        _running_sessions.discard(session_id)


@router.get("/{session_id}/audio/{message_id}")
async def get_message_audio(session_id: str, message_id: str):
    """Generate audio on demand via Chatterbox (HuggingFace) — FREE."""
    import json as jsonlib
    import os
    import re
    import uuid
    import httpx as hx

    result = db.get_supabase().table("forum_messages").select("metadata,content,agent_name").eq("id", message_id).execute()
    if not result.data:
        raise HTTPException(status_code=404, detail="Message not found")

    msg = result.data[0]
    existing_url = msg.get("metadata", {}).get("audio_url")
    if existing_url:
        return {"audio_url": existing_url}

    # Clean text — take first 250 chars of meaningful content
    text = msg["content"]
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\[CHALLENGE:[^\]]+\]", "", text)
    text = re.sub(r"#{1,3}\s*", "", text)
    text = re.sub(r"^(DECLARACIÓN|APOYO|INTERPELACIÓN|RESPUESTA)\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\n+", " ", text).strip()
    # Take first 2 sentences max 250 chars
    sentences = re.split(r"(?<=[.!?])\s+", text)
    text = " ".join(sentences[:2])[:250].strip()

    if len(text) < 5:
        raise HTTPException(status_code=400, detail="Text too short")

    # Voice seeds per agent
    seeds = {"Elena Vásquez": 42, "Marcus Chen": 137, "Sofia Andersen": 256, "Ahmed Al-Rashid": 789, "Dr. Ingrid Hoffmann": 512, "James Okafor": 333, "Laura Martínez": 611, "Moderador": 100, "Integrador": 200}
    seed = seeds.get(msg["agent_name"], 42)

    jwt = os.environ.get("SUPABASE_JWT_KEY", "")
    supa_url = os.environ.get("SUPABASE_URL", "")

    try:
        with hx.Client(timeout=120.0) as client:
            # Step 1: Submit to Chatterbox
            resp = client.post(
                "https://resembleai-chatterbox.hf.space/gradio_api/call/generate_tts_audio",
                headers={"Content-Type": "application/json"},
                json={"data": [text, None, 0.4, 0.9, seed, 0.5, False]},
            )
            if resp.status_code != 200:
                raise HTTPException(status_code=500, detail=f"Chatterbox submit: {resp.status_code}: {resp.text[:200]}")

            event_id = resp.json().get("event_id")
            if not event_id:
                raise HTTPException(status_code=500, detail=f"No event_id: {resp.text[:200]}")

            # Step 2: Stream result (SSE format)
            result_resp = client.get(
                f"https://resembleai-chatterbox.hf.space/gradio_api/call/generate_tts_audio/{event_id}",
                timeout=120.0,
            )

            audio_hf_url = None
            for line in result_resp.text.split("\n"):
                if line.startswith("data: "):
                    raw = line[6:].strip()
                    if raw and raw != "null":
                        try:
                            data = jsonlib.loads(raw)
                            if isinstance(data, list) and len(data) > 0:
                                item = data[0]
                                if isinstance(item, dict):
                                    audio_hf_url = item.get("url") or item.get("path")
                        except (jsonlib.JSONDecodeError, TypeError, IndexError):
                            pass

            if not audio_hf_url:
                raise HTTPException(status_code=500, detail=f"No audio URL. Response: {result_resp.text[:300]}")

            # Ensure full URL
            if audio_hf_url.startswith("/"):
                audio_hf_url = f"https://resembleai-chatterbox.hf.space{audio_hf_url}"
            elif not audio_hf_url.startswith("http"):
                audio_hf_url = f"https://resembleai-chatterbox.hf.space/gradio_api/file={audio_hf_url}"

            # Step 3: Download audio
            audio_resp = client.get(audio_hf_url, timeout=30.0)
            if audio_resp.status_code != 200:
                raise HTTPException(status_code=500, detail=f"Download failed: {audio_resp.status_code}")
            audio_bytes = audio_resp.content

            if len(audio_bytes) < 1000:
                raise HTTPException(status_code=500, detail="Audio too small, likely empty")

            # Step 4: Upload to Supabase Storage
            if not jwt or not supa_url:
                raise HTTPException(status_code=500, detail="Missing Supabase creds")

            fname = f"{session_id}/{uuid.uuid4().hex[:8]}.wav"
            r = client.post(
                f"{supa_url}/storage/v1/object/audio/{fname}",
                headers={"Authorization": f"Bearer {jwt}", "Content-Type": "audio/wav", "x-upsert": "true"},
                content=audio_bytes,
            )
            if r.status_code not in (200, 201):
                raise HTTPException(status_code=500, detail=f"Storage upload: {r.status_code}: {r.text[:200]}")

        audio_url = f"{supa_url}/storage/v1/object/public/audio/{fname}"
        await db.update_message_metadata(message_id, {"audio_url": audio_url})
        return {"audio_url": audio_url}

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"TTS error: {str(e)}")
